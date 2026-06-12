"""
DOCX Exporter — professional Word document with formatted lead table
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Color palette (as RGB tuples)
NAVY  = RGBColor(0x0D, 0x1E, 0x38)
GOLD  = RGBColor(0xD4, 0xAF, 0x37)
LIGHT = RGBColor(0xEF, 0xF3, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x1A, 0x2E)


def _hex(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


def _set_cell_bg(cell, r, g, b):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(r, g, b))
    tcPr.append(shd)


def export_docx(leads: list, output_dir: str, filename: str = None) -> str:
    if not filename:
        filename = f"leads_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = os.path.join(output_dir, filename)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width  = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = section.right_margin = Inches(0.75)
        section.top_margin  = section.bottom_margin = Inches(0.75)

    # ── Title block ───────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("FreelancerX Lead Scraper — Lead Report")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}  ·  Total Leads: {len(leads)}"
    )
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = GOLD

    doc.add_paragraph()

    # ── Summary box ───────────────────────────────────────────────────────
    with_email = sum(1 for l in leads if l.get("email"))
    with_phone = sum(1 for l in leads if l.get("phone"))
    with_site  = sum(1 for l in leads if l.get("website"))

    summary_tbl = doc.add_table(rows=1, cols=3)
    summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_tbl.style = "Table Grid"
    for i, (label, val) in enumerate([("With Email", with_email), ("With Phone", with_phone), ("With Website", with_site)]):
        cell = summary_tbl.rows[0].cells[i]
        _set_cell_bg(cell, 0x0D, 0x1E, 0x38)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{val}\n")
        r1.font.size = Pt(18)
        r1.font.bold = True
        r1.font.color.rgb = GOLD
        r2 = p.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = WHITE

    doc.add_paragraph()

    # ── Lead table ────────────────────────────────────────────────────────
    COLS = ["#", "Business Name", "Email", "Phone", "Website", "City", "Country"]
    FIELDS = ["_idx", "business_name", "email", "phone", "website", "city", "country"]
    COL_WIDTHS = [0.35, 2.0, 2.0, 1.4, 2.2, 1.2, 1.2]

    table = doc.add_table(rows=1 + len(leads), cols=len(COLS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, (col, w) in enumerate(zip(COLS, COL_WIDTHS)):
        cell = hdr_row.cells[i]
        cell.width = Inches(w)
        _set_cell_bg(cell, 0x0D, 0x1E, 0x38)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = GOLD

    # Data rows
    for row_i, lead in enumerate(leads):
        row = table.rows[row_i + 1]
        fill = (0xEF, 0xF3, 0xFA) if row_i % 2 == 0 else (0xFF, 0xFF, 0xFF)

        for col_i, field in enumerate(FIELDS):
            cell = row.cells[col_i]
            cell.width = Inches(COL_WIDTHS[col_i])
            _set_cell_bg(cell, *fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_i > 0 else WD_ALIGN_PARAGRAPH.CENTER

            val = str(row_i + 1) if field == "_idx" else (lead.get(field, "") or "")
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(8)
            run.font.color.rgb = BLACK

    doc.add_paragraph()

    footer_para = doc.add_paragraph("FreelancerX Lead Scraper — freelankarx.github.io")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xCC)

    doc.save(path)
    return path
